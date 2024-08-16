import streamlit as st
import requests
from docx import Document
from io import BytesIO
import time

# User credentials
USERNAME = "Vikram_Lingam"
PASSWORD = "Vikram123"

# Expanded list of IFRS-related keywords and phrases
IFRS_KEYWORDS = [
    "IFRS", "financial statements", "revenue recognition", "lease", "asset",
    "liability", "equity", "financial instrument", "impairment", "fair value",
    "consolidation", "subsidiary", "cash flow", "income statement", "balance sheet",
    "taxation", "deferred tax", "depreciation", "amortization", "goodwill", "provision",
    "acquisition", "business combination", "disclosure", "hedge accounting", "intangible assets",
    "investment property", "joint venture", "minority interest", "non-controlling interest",
    "operating segment", "other comprehensive income", "retained earnings", "share capital",
    "statement of changes in equity", "stock compensation", "conceptual framework", "IAS", "Interim", "Interpretation"
    "Funds", "Rehabilitation", "Liabilities", "Economies"
]

# Function to create a Word document
def create_word_doc(query, background, standard_summary, analysis, conclusion):
    doc = Document()
    
    # Adding heading and sections
    doc.add_heading(query, level=1)
    doc.add_heading('Background', level=2)
    doc.add_paragraph(background)
    doc.add_heading('IFRS Standard Summary', level=2)
    doc.add_paragraph(standard_summary)
    doc.add_heading('Analysis', level=2)
    doc.add_paragraph(analysis)
    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(conclusion)
    
    # Saving to a bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Function to simulate typing effect
def type_text(text, speed=0.05):
    placeholder = st.empty()  # Create an empty placeholder that we will update
    typed_text = ""

    for char in text:
        typed_text += char
        placeholder.markdown(f"<p style='font-family:monospace; font-size:16px'>{typed_text}</p>", unsafe_allow_html=True)
        time.sleep(speed)

# Function to check if the query is IFRS related
def is_ifrs_related(query):
    for keyword in IFRS_KEYWORDS:
        if keyword.lower() in query.lower():
            return True
    return False

# Authentication function
def check_credentials(username, password):
    return username == USERNAME and password == PASSWORD

# Streamlit sidebar for login
st.sidebar.title("Login")
username_input = st.sidebar.text_input("Username")
password_input = st.sidebar.text_input("Password", type="password")
login_button = st.sidebar.button("Login")

# Session state to keep track of login status
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if login_button:
    if check_credentials(username_input, password_input):
        st.session_state.logged_in = True
        st.sidebar.success("Login successful!")
    else:
        st.sidebar.error("Invalid username or password.")

# Main content
if st.session_state.logged_in:
    st.title("IFRS Research Assistant")
    st.write("Author: **Vikram Lingam**")
    st.write("GitHub Repo: [https://github.com/vikramlingam](https://github.com/vikramlingam)")
    st.write("Enter your IFRS-related query to get detailed research insights.")

    # User input for IFRS query
    query = st.text_input("Enter your query:", "")

    if st.button("Search"):
        if query:
            # Check if the query is related to IFRS
            if is_ifrs_related(query):
                # Prepare the API request
                url = "https://api.perplexity.ai/chat/completions"
                payload = {
                    "model": "llama-3.1-sonar-small-128k-online",
                    "messages": [
                        {
                            "role": "system",
                            "content": "Be precise and concise."
                        },
                        {
                            "role": "user",
                            "content": query
                        }
                    ]
                }
                headers = {
                    "accept": "application/json",
                    "content-type": "application/json",
                    "Authorization": f"Bearer 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"
                }
                
                response = requests.post(url, json=payload, headers=headers)

                if response.status_code == 200:
                    results = response.json()
                    content = results['choices'][0]['message']['content'] if 'choices' in results and len(results['choices']) > 0 else "Not Available"
                    
                    # Display the typing effect
                    st.write("### Results:")
                    type_text(content)
                    
                    # Example content splitting into sections (modify as per actual response structure)
                    background = "Background extracted from the content if possible."
                    standard_summary = "IFRS Standard summary."
                    analysis = "Analysis based on the content."
                    conclusion = "Conclusion based on the content."

                    # Option to download the response as a Word document
                    buffer = create_word_doc(query, background, standard_summary, analysis, conclusion)
                    st.download_button(label="Download as Word Document", data=buffer, file_name=f"{query}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
                else:
                    st.write(f"Error: {response.status_code}")
                    st.write(f"Response: {response.content}")
            else:
                # If the query is not IFRS-related
                st.write("Out of IFRS context.")
        else:
            st.write("Please enter a query.")
    
    # Adding a footer with IFRS standards link
    st.write("---")
    st.write("For detailed information about IFRS standards, visit [IFRS Standards](https://www.ifrs.org/issued-standards/list-of-standards/).")
else:
    st.write("Please log in to use the IFRS Research Assistant.")
